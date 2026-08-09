import streamlit as st
import re
import pandas as pd
from io import BytesIO
from pypdf import PdfReader
import docx
import phonenumbers
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# pdfplumber اختياري: لو غير متوفر بأي بيئة تشغيل عند العميل، يرجع تلقائيًا لـ pypdf
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# ---------------------------------------------------------
# 1. القواعد والبيانات الأساسية
# ---------------------------------------------------------
EMAIL_REGEX = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
PHONE_FALLBACK_REGEX = r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}'
EXP_REGEX = r'(\d{1,2})\+?\s*(?:-\s*\d{1,2})?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)'

SKILLS_DATABASE = [
    'Python', 'Java', 'JavaScript', 'C++', 'Machine Learning', 'Deep Learning',
    'Data Analysis', 'Data Science', 'SQL', 'NoSQL', 'NLP', 'Natural Language Processing',
    'Pandas', 'NumPy', 'Arabic NLP', 'Dart', 'Flutter', 'React', 'Node.js',
    'Oracle APEX', 'Relational Databases', 'Arduino', 'ATmega', 'EV3 Robotics',
    'Sensor Integration', 'Scikit-learn', 'TensorFlow', 'PyTorch', 'BERT', 'Word2Vec',
    'TF-IDF', 'BOW', 'CNN', 'RNN', 'Computer Vision', 'Image Processing',
    'Cloud Computing', 'AWS', 'Azure', 'Docker', 'Git', 'HTML', 'CSS',
    'Project Management', 'Excel', 'Power BI', 'Tableau'
]

EDUCATION_KEYWORDS = [
    'Bachelor', 'Master', 'PhD', 'B.Sc', 'M.Sc', 'Diploma',
    'Community College', 'Artificial Intelligence', 'Computer Science',
    'Information Technology', 'Engineering', 'Data Science', 'Business Administration'
]

NAME_LINE_BLOCKLIST = {'resume', 'cv', 'curriculum', 'vitae', 'profile',
                        'contact', 'address', 'summary', 'objective', 'personal'}

# ---------------------------------------------------------
# 2. تنظيف النص
# ---------------------------------------------------------
def clean_extracted_text(text):
    if not text:
        return ""
    cleaned = re.sub(r'(?<=\b[a-zA-Z])\s+(?=[a-zA-Z]\b)', '', text)
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    return cleaned

# ---------------------------------------------------------
# 3. استخراج النصوص من الملفات
# ---------------------------------------------------------
def extract_text_from_pdf(file_bytes):
    text = ""
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            text = ""

    if len(text.strip()) < 30:
        try:
            reader = PdfReader(BytesIO(file_bytes))
            fallback_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    fallback_text += page_text + "\n"
            if len(fallback_text.strip()) > len(text.strip()):
                text = fallback_text
        except Exception:
            pass

    return clean_extracted_text(text)

def extract_text_from_docx(file_bytes):
    doc = docx.Document(BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return clean_extracted_text("\n".join(parts))

def parse_file(uploaded_file):
    bytes_data = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(bytes_data)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(bytes_data)
    elif filename.endswith('.txt'):
        return clean_extracted_text(bytes_data.decode('utf-8', errors='ignore'))
    return ""

# ---------------------------------------------------------
# 4. استخراج الحقول
# ---------------------------------------------------------
def extract_name(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:6]:
        lower = line.lower()
        if any(k in lower for k in NAME_LINE_BLOCKLIST):
            continue
        if re.search(r'@|\d{3,}|http|www|linkedin|github', line, re.IGNORECASE):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w[:1].isalpha()):
            return " ".join(words)
    return "غير محدد"

def extract_email(text):
    match = re.search(EMAIL_REGEX, text)
    return match.group(0) if match else "غير موجود"

def extract_phone(text):
    for match in phonenumbers.PhoneNumberMatcher(text, "OM"):
        return phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
    fallback = re.search(PHONE_FALLBACK_REGEX, text)
    return fallback.group(0).strip() if fallback else "غير موجود"

def extract_experience(text):
    matches = re.findall(EXP_REGEX, text, re.IGNORECASE)
    if matches:
        years = [int(m) for m in matches]
        return f"{max(years)}+ سنوات خبرة"
    return "غير محدد / حديث تخرج"

def parse_resume(text, filename):
    found_skills = [s for s in SKILLS_DATABASE
                     if re.search(r'\b' + re.escape(s) + r'\b', text, re.IGNORECASE)]
    found_education = [e for e in EDUCATION_KEYWORDS
                        if re.search(r'\b' + re.escape(e) + r'\b', text, re.IGNORECASE)]

    return {
        'اسم الملف': filename,
        'اسم': extract_name(text),
        'بريد إلكتروني': extract_email(text),
        'هاتف': extract_phone(text),
        'خبرة': extract_experience(text),
        'مهارات': ", ".join(found_skills) if found_skills else "غير محدد",
        'تعليم': ", ".join(sorted(set(found_education))) if found_education else "غير محدد",
        '_skills_list': found_skills,
        'Raw_Text': text
    }

# ---------------------------------------------------------
# 5. مطابقة الوظائف (تشابه نصي + تطابق مهارات)
# ---------------------------------------------------------
def calculate_match_score(job_description, resume_text, resume_skills):
    if not job_description.strip() or not resume_text.strip():
        return 0.0

    vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
    tfidf_matrix = vectorizer.fit_transform([job_description, resume_text])
    text_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

    jd_skills = [s for s in SKILLS_DATABASE
                 if re.search(r'\b' + re.escape(s) + r'\b', job_description, re.IGNORECASE)]
    overlap = (len(set(jd_skills) & set(resume_skills)) / len(jd_skills)) if jd_skills else 0.0

    final_score = (0.5 * text_score) + (0.5 * overlap)
    return round(final_score * 100, 2)

# ---------------------------------------------------------
# 6. تصدير Excel حقيقي
# ---------------------------------------------------------
def to_excel_bytes(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Resumes')
    return output.getvalue()

# ---------------------------------------------------------
# 7. الواجهة
# ---------------------------------------------------------
st.set_page_config(page_title="محلل السيرة الذاتية الاحترافي", layout="wide")
st.title("📄 محلل السيرة الذاتية بالذكاء الاصطناعي ومطابقة الوظائف")
st.caption("استخراج تلقائي للبيانات من PDF / DOCX / TXT، تصدير منظم، وتقييم توافق ذكي مع الوظيفة.")

st.sidebar.header("🎯 مطابقة الوظيفة (اختياري)")
job_description = st.sidebar.text_area("أدخل وصف الوظيفة (Job Description):", height=250)

uploaded_files = st.file_uploader("قم برفع السير الذاتية (PDF, DOCX, TXT):",
                                    type=['pdf', 'docx', 'txt'], accept_multiple_files=True)

if uploaded_files:
    parsed_results = []
    for file in uploaded_files:
        text = parse_file(file)
        data = parse_resume(text, file.name)
        if job_description.strip():
            data['نتيجة المباراة (%)'] = calculate_match_score(job_description, text, data['_skills_list'])
        else:
            data['نتيجة المباراة (%)'] = None
        del data['_skills_list']
        parsed_results.append(data)

    df = pd.DataFrame(parsed_results)

    if job_description.strip():
        df = df.sort_values(by='نتيجة المباراة (%)', ascending=False).reset_index(drop=True)
        df['نتيجة المباراة (%)'] = df['نتيجة المباراة (%)'].apply(lambda x: f"{x}%")
    else:
        df['نتيجة المباراة (%)'] = "غير متوفر"

    st.subheader("📊 البيانات المستخرجة")
    display_df = df.drop(columns=['Raw_Text'])
    st.dataframe(display_df, use_container_width=True)

    st.subheader("📥 تصدير النتائج")
    col1, col2, col3 = st.columns(3)
    export_df = df.drop(columns=['Raw_Text'])

    col1.download_button("⬇️ تحميل JSON",
                          data=export_df.to_json(orient="records", indent=4, force_ascii=False),
                          file_name="parsed_resumes.json", mime="application/json")

    col2.download_button("⬇️ تحميل Excel (.xlsx)",
                          data=to_excel_bytes(export_df),
                          file_name="parsed_resumes.xlsx",
                          mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    col3.download_button("⬇️ تحميل CSV",
                          data=export_df.to_csv(index=False).encode('utf-8-sig'),
                          file_name="parsed_resumes.csv", mime="text/csv")
